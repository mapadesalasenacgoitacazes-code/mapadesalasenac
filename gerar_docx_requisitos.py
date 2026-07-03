"""
Gera o Documento de Requisitos Unificado v3.0 em formato Word (.docx)
com comentarios da Thais Oliveira nas partes que ela acrescentou/alterou.
"""
import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from lxml import etree

AZUL = RGBColor(0x00, 0x4C, 0x94)
LARANJA = RGBColor(0xF7, 0x94, 0x1D)
BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
CINZA = RGBColor(0x66, 0x66, 0x66)
PRETO = RGBColor(0x1A, 0x1A, 0x1A)
VERDE_ESCURO = RGBColor(0x16, 0x65, 0x34)
VERMELHO = RGBColor(0xC6, 0x28, 0x28)

cid = [0]


def next_cid():
    cid[0] += 1
    return str(cid[0])


def add_tracked_comment(paragraph, text, all_comments, author="Thais Oliveira", initials="TO"):
    """Add a Word comment and track it."""
    c = next_cid()
    _add_comment_to_paragraph(paragraph, c, author, initials, text)
    all_comments.append((c, author, initials, text))


def _add_comment_to_paragraph(paragraph, cid, author, initials, text):
    """Low-level comment insertion via OOXML."""
    # Create commentRangeStart
    comment_range_start = OxmlElement('w:commentRangeStart')
    comment_range_start.set(qn('w:id'), cid)

    # Create commentRangeEnd
    comment_range_end = OxmlElement('w:commentRangeEnd')
    comment_range_end.set(qn('w:id'), cid)

    # Create commentReference run
    comment_ref_run = OxmlElement('w:r')
    comment_ref_rpr = OxmlElement('w:rPr')
    comment_ref_style = OxmlElement('w:rStyle')
    comment_ref_style.set(qn('w:val'), 'CommentReference')
    comment_ref_rpr.append(comment_ref_style)
    comment_ref_run.append(comment_ref_rpr)
    comment_reference = OxmlElement('w:commentReference')
    comment_reference.set(qn('w:id'), cid)
    comment_ref_run.append(comment_reference)

    p_elem = paragraph._element
    if len(p_elem) > 0:
        p_elem.insert(0, comment_range_start)
    else:
        p_elem.append(comment_range_start)
    p_elem.append(comment_range_end)
    p_elem.append(comment_ref_run)

    return cid, author, initials, text


def finalize_comments(doc, comments_data):
    """Create the comments part in the document."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.opc.part import Part
    from docx.opc.packuri import PackURI

    comments_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
    comments_xml += '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
    comments_xml += 'xmlns:r="http://schemas.openxmlformats.org/officeDocument2006/relationships">'

    for cid, author, initials, text in comments_data:
        date_str = datetime.datetime.now().strftime("%Y-%m-%dT%H:%M:%S") + "Z"
        comments_xml += f'<w:comment w:id="{cid}" w:author="{author}" w:date="{date_str}" w:initials="{initials}">'
        comments_xml += '<w:p>'
        comments_xml += '<w:pPr><w:pStyle w:val="CommentText"/></w:pPr>'
        comments_xml += '<w:r><w:rPr><w:rStyle w:val="CommentReference"/></w:rPr>'
        comments_xml += f'<w:annotationRef/></w:r>'
        comments_xml += f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r>'
        comments_xml += '</w:p>'
        comments_xml += '</w:comment>'

    comments_xml += '</w:comments>'

    comments_part = Part(
        PackURI('/word/comments.xml'),
        'application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml',
        comments_xml.encode('utf-8'),
        doc.part.package
    )

    doc.part.relate_to(
        comments_part,
        'http://schemas.openxmlformats.org/officeDocument2006/relationships/comments'
    )


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_styled(doc, text, level=1, color=AZUL):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=11, color=PRETO, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, bold_prefix="", level=0):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
    else:
        if p.runs:
            p.runs[0].font.size = Pt(11)
        else:
            r = p.add_run(text)
            r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table_row(table, cells_data, header=False):
    row = table.add_row()
    for i, (text, bold) in enumerate(cells_data):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        if header:
            set_cell_shading(cell, "004C94")
            run.font.color.rgb = BRANCO
            run.bold = True
        elif bold:
            run.bold = True
    return row


def add_req_box(doc, req_id, title, moscow_pt, description, rules, comments_list, actor=""):
    p_title = doc.add_paragraph()
    r = p_title.add_run(f"{req_id} — {title}")
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = AZUL

    r2 = p_title.add_run(f"  [{moscow_pt}]")
    r2.bold = True
    r2.font.size = Pt(10)
    if "Obrigatorio" in moscow_pt:
        r2.font.color.rgb = VERMELHO
    elif "Importante" in moscow_pt:
        r2.font.color.rgb = LARANJA
    elif "Desejavel" in moscow_pt:
        r2.font.color.rgb = VERDE_ESCURO
    else:
        r2.font.color.rgb = CINZA

    p_desc = add_para(doc, description, size=11)

    if rules:
        p_rules = add_para(doc, f"Regras: {rules}", italic=True, size=10, color=CINZA)

    doc.add_paragraph()
    return p_title, p_desc


def main():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = PRETO

    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    all_comments = []

    # =========================================================
    # CAPA / CABECALHO
    # =========================================================
    title = doc.add_heading('Documento de Requisitos de Software', level=0)
    for run in title.runs:
        run.font.color.rgb = AZUL
        run.font.size = Pt(24)

    subtitle = doc.add_paragraph()
    r = subtitle.add_run('Sistema de Gerenciamento e Mapeamento de Ambientes')
    r.font.size = Pt(14)
    r.font.color.rgb = LARANJA
    r.bold = True
    subtitle.add_run('\n')
    r2 = subtitle.add_run('Mapa de Salas — SENAC Minas BH')
    r2.font.size = Pt(12)
    r2.font.color.rgb = CINZA

    # Metadata table
    meta_table = doc.add_table(rows=0, cols=4)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        [("Versao:", True), ("3.0 (Unificada)", False), ("Data:", True), ("01/07/2026", False)],
        [("Projeto:", True), ("Mapa de Salas SENAC BH", False), ("Equipe:", True), ("Thais Oliveira e equipe", False)],
        [("Professor:", True), ("Cleiton de Jesus Pereira", False), ("Curso:", True), ("Tec. Des. Sistemas - UC 7", False)],
        [("Hospedagem MVP:", True), ("InfinityFree", False), ("Stack:", True), ("PHP 8 + MySQL 8 + Bootstrap 5", False)],
    ]
    for row_data in meta_data:
        row = meta_table.add_row()
        for i, (text, bold) in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.bold = bold
            if bold:
                run.font.color.rgb = AZUL

    doc.add_paragraph()

    # =========================================================
    # 1. OBJETIVO
    # =========================================================
    add_heading_styled(doc, "1. Objetivo da Etapa", level=1)
    add_para(doc, "O objetivo desta fase de Coleta e Analise de Requisitos e entender profundamente como funciona a alocacao de salas de cursos no SENAC Minas (unidade BH), mapear as dores do processo atual (planilha Excel com formato Gantt anual) e transformar as necessidades da equipe administrativa em requisitos funcionais e nao funcionais claros, mensuraveis e testaveis para o desenvolvimento do sistema Mapa de Salas.")
    add_para(doc, "O sistema substituira a planilha atualmente utilizada, que apresenta problemas reais: erros de formula (#N/D), impossibilidade de filtrar por turma, risco de conflitos detectados apenas visualmente e limitacao de visualizacao por ano.")

    # =========================================================
    # 2. PAPEIS
    # =========================================================
    add_heading_styled(doc, "2. Papeis e Responsabilidades", level=1)
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for i, txt in enumerate(["Papel", "Responsabilidade"]):
        hdr.cells[i].text = txt
        set_cell_shading(hdr.cells[i], "004C94")
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.bold = True
                r.font.size = Pt(10)

    papeis = [
        ("Analista de Requisitos", "Conduzir entrevistas, documentar e validar requisitos. Responsavel: Thais Oliveira."),
        ("Stakeholder (Cliente)", "Supervisao / Coordenacao Pedagogica do SENAC BH. Valida requisitos e prioriza funcionalidades."),
        ("Usuarios-chave", "Equipe de Suporte (reserva salas) e Docentes (consultam e reportam problemas)."),
        ("Time Tecnico", "Equipe de desenvolvimento (alunos). Analise de viabilidade e implementacao."),
    ]
    for papel, resp in papeis:
        row = t.add_row()
        row.cells[0].text = papel
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = resp
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================
    # 3. TECNICAS
    # =========================================================
    add_heading_styled(doc, "3. Tecnicas de Coleta Utilizadas", level=1)
    tecnicas = [
        ("Entrevista Individual (1:1): ", "Realizada com a supervisao/cliente para capturar dores, regras de negocio e necessidades especificas."),
        ("Observacao Direta: ", "Analise da planilha Excel (fotos projetadas em sala), identificando campos, erros e fluxos."),
        ("Analise de Documentos: ", "Estudo da planilha 'Mapa de Salas — ANO 2026' com dados reais."),
        ("Prototipacao (Media Fidelidade): ", "Wireframes de 7 telas com paleta oficial SENAC para validacao."),
    ]
    for prefix, text in tecnicas:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # =========================================================
    # 4. CONTEXTO
    # =========================================================
    add_heading_styled(doc, "4. Contexto e Visao Geral do Projeto", level=1)

    add_heading_styled(doc, "4.1. Problema / Oportunidade", level=2)
    add_para(doc, "A equipe de supervisao do SENAC BH gerencia a alocacao de mais de 20 turmas em diversas salas usando uma planilha Excel com formato Gantt anual. A planilha apresenta problemas reais: erros de formula (#N/D), impossibilidade de filtrar por turma, conflitos detectados apenas visualmente, e limitacao por ano.")

    add_heading_styled(doc, "4.2. Objetivos de Negocio", level=2)
    objetivos = [
        "Eliminar conflitos de sala com validacao automatica.",
        "Consulta instantanea do historico de salas por turma.",
        "Calculo automatico de progresso e ajuste de data de termino.",
        "Gestao centralizada com controle de acesso por perfil.",
        "4 modos de visualizacao (Dia, Semana, Mes, Ano).",
    ]
    for o in objetivos:
        add_bullet(doc, o)

    add_heading_styled(doc, "4.3. Premissas", level=2)
    premissas = [
        "A equipe cadastra as salas livremente (quantidade, tipo, andar, capacidade).",
        "Cada turma tem dias da semana fixos, cadastrados pela equipe.",
        "Cada unidade define seus proprios feriados e recessos.",
        "Sabado e letivo apenas para algumas turmas.",
        "Reserva por turno (Manha/Tarde/Noite), nao por horario livre.",
        "Uma turma pode ter mais de um professor vinculado.",
        "PHP + MySQL como stack obrigatoria (definido pelo professor).",
    ]
    for p in premissas:
        add_bullet(doc, p)

    add_heading_styled(doc, "4.4. Restricoes", level=2)
    restricoes = [
        ("Stack: ", "PHP 8 + MySQL 8, MVC puro sem framework (definido pelo professor)."),
        ("Hospedagem MVP: ", "InfinityFree (PHP + MySQL gratuito, sem SSH, SSL via Let's Encrypt)."),
        ("Frontend: ", "HTML5 + CSS3 + JavaScript + Bootstrap 5."),
        ("Prazo: ", "Projeto academico, semestre letivo 2026.2."),
        ("Navegador: ", "Chrome/Edge. Sem necessidade de app mobile nativo na v1."),
    ]
    for prefix, text in restricoes:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # =========================================================
    # 5. ATORES
    # =========================================================
    add_heading_styled(doc, "5. Atores do Sistema (Perfis de Acesso)", level=1)
    p_atores = add_para(doc, "4 perfis definidos a partir das entrevistas com a cliente e do roteiro preenchido em sala:")

    add_tracked_comment(p_atores,
        "MINHA CONTRIBUICAO: Corrigi os perfis do documento do grupo (Master/Admin/Padrao) para os nomes reais confirmados pela cliente na entrevista: Admin, Supervisao, Suporte e Docente. O grupo usou terminologia generica que nao corresponde ao vocabulario do SENAC.",
        all_comments)

    t = doc.add_table(rows=1, cols=3)
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for i, txt in enumerate(["Perfil", "Quem", "Permissoes"]):
        hdr.cells[i].text = txt
        set_cell_shading(hdr.cells[i], "004C94")
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.bold = True
                r.font.size = Pt(10)

    atores = [
        ("Admin", "Futuro responsavel eleito pela unidade", "Gerenciar usuarios, configurar sistema, CRUD de salas, supervisionar tudo. Acesso total."),
        ("Supervisao", "Coordenacao Pedagogica", "Visualizar tudo, dashboard, relatorios, log de auditoria. Decisoes estrategicas."),
        ("Suporte", "Equipe operacional (Secretaria)", "CRUD reservas, alterar status, cadastrar turmas/cursos, gerenciar feriados."),
        ("Docente", "Professores", "Visualizar ocupacao, consultar reservas e turmas. Somente leitura."),
    ]
    for perfil, quem, perm in atores:
        row = t.add_row()
        row.cells[0].text = perfil
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = quem
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[2].text = perm
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)

    p_nota = add_para(doc, 'Nota da entrevista: A cliente informou que "no futuro eles vao eleger uma pessoa pra admin do sistema". O perfil Admin existe no banco mas pode nao ser utilizado na v1.', italic=True, size=10, color=CINZA)

    doc.add_paragraph()

    # =========================================================
    # 6. REQUISITOS FUNCIONAIS
    # =========================================================
    add_heading_styled(doc, "6. Requisitos Funcionais", level=1)

    # RF01
    p1, _ = add_req_box(doc, "RF01", "Autenticacao e Controle de Acesso",
        "Must — Obrigatorio",
        "O sistema deve permitir login com e-mail e senha, com 4 perfis diferenciados: Admin, Supervisao, Suporte e Docente. Cada perfil define permissoes especificas de leitura e escrita.",
        "Senha min. 8 caracteres (letras + numeros + simbolos). E-mail unico. Sessao expira em 30 min. password_hash() (bcrypt). session_regenerate_id(). Token CSRF em formularios. Atores: Todos (login), Admin (gerenciar usuarios).",
        all_comments)

    # RF02
    p2, pd2 = add_req_box(doc, "RF02", "Cadastro de Salas e Ambientes",
        "Must — Obrigatorio",
        "O sistema deve permitir cadastrar salas informando: numero/nome, tipo (Sala de Aula, Laboratorio, Auditorio, Cozinha, Especial), capacidade (vagas), andar/pavimento (Terreo, 1o andar, 2o andar, etc.), recursos disponiveis e observacoes.",
        "Numero unico por unidade. Tipo, capacidade e andar obrigatorios. Status inicial: Disponivel. Admin e Suporte podem cadastrar. Banco: tabelas sala + recurso_sala. UNIQUE(id_unidade, numero).",
        all_comments)
    add_tracked_comment(pd2,
        "MINHA CONTRIBUICAO: Acrescentei o campo ANDAR como obrigatorio apos confirmar com a cliente que o predio tem salas em andares diferentes e que o numero do andar e importante. O documento do grupo nao incluia esse campo.",
        all_comments)

    # RF03
    p3, pd3 = add_req_box(doc, "RF03", "Cadastro de Cursos",
        "Should — Importante",
        "O sistema deve permitir cadastrar cursos com: nome, tipo de curso (Cursos Ageis, Graduacao, Pos Graduacao, Tecnico), programa (Todos os cursos, Senac+, PSG), area de conhecimento (23 areas padrao SENAC) e carga horaria total.",
        "Nome, tipo, area e carga horaria obrigatorios. ENUMs baseados no catalogo oficial SENAC Minas. Programa padrao: 'Todos'. 23 areas: Gestao, Comercio, Comunicacao, Artes, Design, Saude, TI, Idioma, Educacional, Lazer, Moda, Beleza, Turismo, Social, Producao de Alimentos, Seguranca, Hospedagem, Gastronomia, Eventos, Asseio e Conservacao, Meio Ambiente e Saude, Meio Ambiente e Recursos Naturais, Games. Banco: tabela curso.",
        all_comments)
    add_tracked_comment(pd3,
        "MINHA CONTRIBUICAO: Este requisito inteiro foi adicionado por mim. O documento do grupo nao diferenciava tipos de curso, areas de conhecimento nem programa. Pesquisei no site oficial do SENAC Minas para levantar as 23 areas e os 4 tipos de curso, alem dos programas Senac+ e PSG.",
        all_comments)

    # RF04
    p4, pd4 = add_req_box(doc, "RF04", "Cadastro e Edicao de Turmas",
        "Must — Obrigatorio",
        "O sistema deve permitir cadastrar e editar turmas com: codigo (formato 006.2026.XXXX), curso vinculado, um ou mais professores (com indicacao do principal), turno (Manha/Tarde/Noite), data de inicio, data de termino prevista, dias da semana fixos (Ex: Seg a Sex, somente Sabado, Seg/Qua/Sex), descricao e observacoes.",
        "Codigo unico. Curso e pelo menos 1 professor obrigatorios. Data de termino posterior ao inicio. Descricao e observacao opcionais, disponiveis no cadastro e na edicao. Banco: turma + professor_turma (N:N). UNIQUE(id_turma, id_usuario).",
        all_comments)
    add_tracked_comment(pd4,
        "MINHA CONTRIBUICAO: Corrigi a reserva de 'horario livre' para TURNO (Manha/Tarde/Noite) conforme a cliente usa. Adicionei: multiplos professores (N:N), campos descricao e observacao, e possibilidade de edicao a qualquer momento. O documento do grupo usava 'horario de inicio/termino/duracao' e tratava professor como 1:1.",
        all_comments)

    # RF05
    p5, pd5 = add_req_box(doc, "RF05", "Cadastro de Feriados e Recessos",
        "Should — Importante",
        "O sistema deve permitir que cada unidade cadastre seus proprios feriados e recessos (datas bloqueadas), impedindo alocacao nesses dias e recalculando automaticamente a data de termino dos cursos afetados.",
        "Tipos: Nacional, Estadual, Municipal, Recesso. Data e descricao obrigatorios. Alertar sobre reservas existentes. Cada unidade define seus proprios feriados independentemente. Banco: tabela feriado. UNIQUE(id_unidade, data).",
        all_comments)
    add_tracked_comment(pd5,
        "MINHA CONTRIBUICAO: Adicionei este requisito completo. A cliente confirmou que cada unidade define seus proprios feriados/recessos. O documento do grupo nao menciona feriados em nenhum lugar.",
        all_comments)

    # RF06
    add_req_box(doc, "RF06", "Criar Reserva de Sala",
        "Must — Obrigatorio",
        "O sistema deve permitir alocar uma turma em uma sala para uma data e turno especificos (Manha, Tarde ou Noite). O sistema deve impedir automaticamente reservas conflitantes (mesma sala + mesmo turno + mesmo dia).",
        "Se houver conflito, exibir mensagem clara indicando qual turma ja ocupa e sugerir salas disponiveis. Sala em manutencao nao aceita reserva. Admin e Suporte criam. Banco: reserva com UNIQUE(id_sala, data, turno).",
        all_comments)

    # RF07
    p7, pd7 = add_req_box(doc, "RF07", "Reserva Recorrente (Automatica)",
        "Should — Importante",
        "Ao cadastrar uma turma, o sistema deve gerar automaticamente as reservas para todos os dias do curso com base no padrao de dias da semana, turno e sala, excluindo feriados e recessos.",
        "Gerar apenas nos dias do padrao da turma. Pular feriados/recessos. Alertar conflitos antes de confirmar a geracao em lote.",
        all_comments)

    # RF08
    add_req_box(doc, "RF08", "Desocupar / Realocar Turma",
        "Could — Desejavel",
        "O sistema deve permitir mover uma turma de uma sala para outra (troca pontual ou permanente), mantendo o historico da sala anterior.",
        "Sala de destino disponivel. Motivo registrado. Historico preservado.",
        all_comments)

    # RF09
    p9, pd9 = add_req_box(doc, "RF09", "Registrar Dia sem Aula",
        "Should — Importante",
        "Quando um dia de aula e cancelado (falta do professor, evento, etc.), o sistema deve registrar a ausencia e automaticamente acrescentar +1 dia util ao termino do curso.",
        "O dia adicionado pula feriados, recessos e dias fora do padrao da turma. Motivo obrigatorio. Progresso recalculado. Banco: tabela dia_cancelado. UNIQUE(id_turma, data).",
        all_comments)
    add_tracked_comment(pd9,
        "MINHA CONTRIBUICAO: Levantei esta regra direto da entrevista com a cliente: 'quando tem uma data que nao teve aula, aumenta um dia no final do curso'. Criei a tabela dia_cancelado no schema e a logica de recalculo. O documento do grupo nao menciona esta regra.",
        all_comments)

    # RF10
    add_req_box(doc, "RF10", "Alterar Status da Sala",
        "Must — Obrigatorio",
        "O sistema deve permitir mudar o status de uma sala para: Disponivel, Ocupada ou Manutencao.",
        "Manutencao lista reservas futuras e alerta. Novas reservas bloqueadas. Apenas Suporte pode alterar.",
        all_comments)

    # RF11
    p11, pd11 = add_req_box(doc, "RF11", "Visualizacao Geral (Painel / Calendario)",
        "Must — Obrigatorio",
        "O sistema deve exibir um painel com 4 modos de visualizacao: Dia (detalhe de um unico dia), Semana (7 dias), Mes (calendario mensal) e Ano (visao anual). Salas nas linhas, dias/periodos nas colunas. Cada reserva mostra codigo da turma e turno (cor diferenciada). Feriados e recessos indicados visualmente.",
        "Alternar entre 4 modos a qualquer momento. Visao padrao: Semana. Layout semelhante ao Gantt da planilha. Salas em manutencao em vermelho. Todos visualizam.",
        all_comments)
    add_tracked_comment(pd11,
        "MINHA IDEIA: Propus os 4 modos de visualizacao (Dia/Semana/Mes/Ano) inspirada no Google Agenda. O documento do grupo mencionava apenas 'tabela, grade ou calendario' sem especificar os modos. Acredito que facilita a transicao da planilha Gantt.",
        all_comments)

    # RF12
    p12, pd12 = add_req_box(doc, "RF12", "Filtros Diversos",
        "Must — Obrigatorio",
        "O sistema deve permitir filtrar reservas e turmas por: data, turno (Manha/Tarde/Noite), turma, professor(es), tipo de sala, andar, curso, tipo de curso, area de conhecimento, programa, status da sala, dias da semana e faixa de progresso.",
        "Filtros combinaveis (ex: turno Tarde + Laboratorio + Gastronomia). Disponiveis no painel e na listagem de turmas. Todos os perfis. Tipo de curso e area seguem catalogo oficial SENAC.",
        all_comments)
    add_tracked_comment(pd12,
        "MINHA CONTRIBUICAO: Expandi os filtros para incluir tipo de curso, area de conhecimento, programa, andar e faixa de progresso. Pesquisei os filtros disponiveis no site oficial do SENAC Minas para usar como referencia. O grupo listava apenas filtros basicos.",
        all_comments)

    # RF13
    p13, pd13 = add_req_box(doc, "RF13", "Historico de Salas por Turma",
        "Should — Importante",
        "Ao consultar uma turma pelo codigo, o sistema deve exibir todas as salas pelas quais ela passou ao longo do curso, com datas, tipo de sala e motivo da troca.",
        "Pedido direto da cliente ('e a informacao que mais demora pra encontrar na planilha'). Funciona para turmas multi-sala (ex: Gastronomia). Todos consultam. Banco: view vw_historico_salas_turma.",
        all_comments)
    add_tracked_comment(pd13,
        "MINHA CONTRIBUICAO: Identifiquei esta necessidade direto da entrevista. A cliente disse que a informacao mais dificil de encontrar na planilha e 'por quais salas uma turma passou'. Criei a view vw_historico_salas_turma no banco. O grupo nao inclui esse requisito.",
        all_comments)

    # RF14
    add_req_box(doc, "RF14", "Calculo e Exibicao do Progresso",
        "Should — Importante",
        "O sistema deve calcular automaticamente o percentual de progresso de cada turma: (dias de aula realizados / total de dias previstos) x 100%. Exibir com barra de progresso visual.",
        "Dias cancelados nao contam como realizados. Feriados e recessos nao contam como previstos. Atualizado automaticamente. Banco: view vw_progresso_turma.",
        all_comments)

    # RF15
    add_req_box(doc, "RF15", "Gerenciar Docentes e Turmas",
        "Should — Importante",
        "O sistema deve permitir vincular e desvincular professores de turmas, indicando professor principal. Uma turma pode ter varios professores (N:N).",
        "Pelo menos um professor obrigatorio por turma. Banco: tabela professor_turma + view vw_professores_turma.",
        all_comments)

    # RF16
    add_req_box(doc, "RF16", "Log de Alteracoes (Auditoria)",
        "Could — Desejavel",
        "O sistema deve registrar automaticamente quem fez cada alteracao (criacao, edicao, exclusao de reserva, troca de status), com data/hora e descricao da acao.",
        "Log nao pode ser editado. Apenas Supervisao/Admin. Registrar: usuario, acao, tabela, data/hora. Banco: tabela log_alteracao.",
        all_comments)

    # RF17
    add_req_box(doc, "RF17", "Dashboard de Ocupacao",
        "Could — Desejavel",
        "Painel gerencial: salas livres/ocupadas no momento, taxa de ocupacao por turno, salas em manutencao, turmas com maior demanda.",
        "Apenas Supervisao/Admin. Dados atualizados ao carregar. Bootstrap cards + graficos simples. Banco: view vw_ocupacao_hoje.",
        all_comments)

    # =========================================================
    # 7. REQUISITOS NAO FUNCIONAIS
    # =========================================================
    add_heading_styled(doc, "7. Requisitos Nao Funcionais", level=1)

    # Hosting strategy box
    p_host = add_para(doc, "", size=11)
    r1 = p_host.add_run("Estrategia de Hospedagem")
    r1.bold = True
    r1.font.color.rgb = AZUL
    r1.font.size = Pt(12)

    add_para(doc, "MVP (Apresentacao ao cliente): InfinityFree — hospedagem gratuita com PHP + MySQL, ideal para demonstrar o sistema ao cliente sem custo. Sem SSH, com SSL gratuito via Let's Encrypt. Limitacoes: sem cron jobs avancados, sem monitoramento, sem SLA garantido.", size=11)

    p_prod = add_para(doc, "Producao (Boas Praticas — se aprovado pelo cliente):", bold=True, size=11)
    add_tracked_comment(p_prod,
        "MINHA IDEIA: Propus separar a hospedagem em duas fases: InfinityFree como MVP para apresentacao e validacao, e opcoes profissionais caso o SENAC queira adotar o sistema em producao. Assim o projeto funciona dentro do orcamento academico mas ja tem um caminho de evolucao.",
        all_comments)

    producao_opcoes = [
        ("Hostinger / HostGator (Shared): ", "~R$10-15/mes. PHP + MySQL, SSL, painel cPanel, backup automatico. Melhor custo-beneficio para producao real."),
        ("VPS (DigitalOcean / Contabo): ", "~R$25-50/mes. Controle total com SSH, cron jobs, Let's Encrypt. Requer conhecimento de sysadmin."),
        ("Servidor interno SENAC: ", "Se a instituicao disponibilizar. Sem custo recorrente, controle total, mas depende de infra interna."),
    ]
    for prefix, text in producao_opcoes:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # App Mobile
    p_mobile = add_para(doc, "", size=11)
    r_mob = p_mobile.add_run("Versao Mobile (Opcional — Futuro)")
    r_mob.bold = True
    r_mob.font.color.rgb = AZUL
    r_mob.font.size = Pt(12)

    add_tracked_comment(p_mobile,
        "MINHA IDEIA: Incluir a opcao de app mobile como evolucao futura. Nao e obrigatoria para o MVP, mas mostra visao de produto ao cliente. A interface web responsiva ja atende tablets e celulares via navegador.",
        all_comments)

    mobile_items = [
        ("PWA (Progressive Web App): ", "Abordagem recomendada. Transforma o sistema web em 'app instalavel' no celular sem precisar de loja (Play Store/App Store). Funciona offline para consultas. Custo zero de desenvolvimento adicional — apenas manifest.json + service worker."),
        ("App Nativo (React Native / Flutter): ", "Maior investimento. So faz sentido se houver necessidade de notificacoes push, GPS ou acesso a camera. Fora do escopo do curso tecnico atual."),
        ("Responsivo (Bootstrap 5): ", "Ja incluido no MVP. A interface se adapta a celulares e tablets automaticamente. Atende 90% dos casos de uso mobile sem desenvolvimento extra."),
    ]
    for prefix, text in mobile_items:
        add_bullet(doc, text, bold_prefix=prefix)

    doc.add_paragraph()

    # RNF01
    add_heading_styled(doc, "RNF01 — Desempenho", level=2)
    rnf01 = [
        "Painel/calendario carrega em menos de 3 segundos em conexao normal.",
        "Indicador de carregamento (spinner) para consultas acima de 2 segundos.",
        "Suportar ate 20 usuarios simultaneos (volume real da equipe SENAC BH).",
        "Feedback visual durante operacoes demoradas (botao desabilitado, loading).",
    ]
    for item in rnf01:
        add_bullet(doc, item)

    # RNF02
    add_heading_styled(doc, "RNF02 — Seguranca", level=2)
    rnf02 = [
        "password_hash() com bcrypt para armazenamento de senhas.",
        "PDO com prepared statements em todas as consultas (prevencao SQL Injection).",
        "htmlspecialchars() em toda saida HTML (prevencao XSS).",
        "session_regenerate_id() apos login bem-sucedido.",
        "Token CSRF em todos os formularios POST.",
        "RBAC (Role-Based Access Control) com validacao em cada operacao.",
        "HTTPS via SSL gratuito (Let's Encrypt no InfinityFree).",
        "Conformidade basica com a LGPD: dados pessoais apenas quando necessario.",
    ]
    for item in rnf02:
        add_bullet(doc, item)

    # RNF03
    add_heading_styled(doc, "RNF03 — Usabilidade", level=2)
    rnf03 = [
        "Responsivo para computador e tablet (Bootstrap 5 grid system).",
        "Layout do calendario semelhante ao Gantt da planilha para facilitar a transicao.",
        "Reserva criada em menos de 2 minutos por usuario novo sem treinamento.",
        "Identidade visual SENAC: Azul #004C94, Laranja #F7941D, Laranja Claro #FDC180.",
        "Mensagens de erro claras, sem detalhes tecnicos (ex: 'Sala ja reservada', nao 'UNIQUE constraint violation').",
    ]
    for item in rnf03:
        add_bullet(doc, item)

    # RNF04
    add_heading_styled(doc, "RNF04 — Disponibilidade", level=2)
    rnf04 = [
        "Disponivel durante horario de funcionamento (7h-22h, segunda a sabado).",
        "MVP (InfinityFree): sem SLA garantido, dependente da disponibilidade do servico gratuito.",
        "Producao: se migrar para hosting pago, disponibilidade de 99% e viavel.",
        "Backup: exportacao manual via phpMyAdmin (MVP) ou automatico via cron + mysqldump (producao).",
    ]
    for item in rnf04:
        add_bullet(doc, item)

    # RNF05
    add_heading_styled(doc, "RNF05 — Plataforma e Stack", level=2)
    rnf05 = [
        "Backend: PHP 8+ com arquitetura MVC pura (sem framework — definido pelo professor).",
        "Banco: MySQL 8 com charset utf8mb4.",
        "Frontend: HTML5 + CSS3 + JavaScript + Bootstrap 5.",
        "Servidor dev: Apache (XAMPP ou Laragon).",
        "Servidor MVP: InfinityFree (Apache + PHP + MySQL gratuito, upload via FTP).",
        "Navegadores: Chrome e Edge (Chromium-based).",
    ]
    for item in rnf05:
        add_bullet(doc, item)

    # RNF06 — MANUTENIBILIDADE (Senior Developer perspective)
    add_heading_styled(doc, "RNF06 — Manutenibilidade", level=2)
    p_manut = add_para(doc, "Consideracoes tecnicas do ponto de vista de um desenvolvedor senior:", bold=True, size=11, color=AZUL)

    add_tracked_comment(p_manut,
        "MINHAS CONSIDERACOES: Escrevi estas recomendacoes de manutenibilidade com base nas boas praticas que estudei. Gostaria que a equipe avaliasse quais dessas praticas conseguimos implementar dentro do prazo do projeto. Algumas podem ficar como meta para a v2.0.",
        all_comments)

    add_para(doc, "Arquitetura e Organizacao do Codigo", bold=True, size=11)
    manut_arq = [
        "MVC rigoroso: toda query SQL fica nos Models, nunca em Controllers ou Views. Controllers apenas recebem requisicoes, chamam Models e redirecionam. Views apenas renderizam HTML com dados ja prontos.",
        "Um arquivo por classe/model. Nomenclatura consistente: SalaModel.php, ReservaController.php, turmas/cadastro.php (view).",
        "Arquivo config/database.php centraliza a conexao PDO. Nenhum outro arquivo cria conexao direta.",
        "Autoload simples via spl_autoload_register() ou includes organizados — sem Composer no MVP, mas preparado para adotar no futuro.",
    ]
    for item in manut_arq:
        add_bullet(doc, item)

    add_para(doc, "Padroes de Codigo", bold=True, size=11)
    manut_pad = [
        "snake_case para variaveis e funcoes PHP ($data_inicio, buscar_salas_disponiveis()).",
        "PascalCase para nomes de classes (SalaModel, ReservaController).",
        "Constantes em UPPER_SNAKE_CASE (MAX_USUARIOS_SIMULTANEOS, TURNOS_VALIDOS).",
        "Indentacao com 4 espacos. Sem mistura de tabs e espacos.",
        "Funcoes curtas: se uma funcao ultrapassa 40 linhas, provavelmente deve ser dividida.",
    ]
    for item in manut_pad:
        add_bullet(doc, item)

    add_para(doc, "Versionamento e Colaboracao", bold=True, size=11)
    manut_git = [
        "Git com commits atomicos e mensagens descritivas em portugues (ex: 'Adicionar validacao de conflito na reserva').",
        "Branch principal: main. Feature branches para cada funcionalidade (feature/cadastro-salas, fix/conflito-turno).",
        "Pull Requests com descricao do que muda e por que. Pelo menos 1 review antes do merge.",
        "Nunca commitar .env, senhas ou credenciais. .gitignore configurado desde o inicio.",
    ]
    for item in manut_git:
        add_bullet(doc, item)

    add_para(doc, "Tratamento de Erros e Logs", bold=True, size=11)
    manut_err = [
        "Em producao: display_errors = Off, log_errors = On. Erros vao para arquivo de log, nunca para a tela do usuario.",
        "Mensagens de erro amigaveis para o usuario. Detalhes tecnicos apenas no log do servidor.",
        "Tabela log_alteracao para auditoria de acoes do usuario (RF16).",
        "Try/catch em operacoes criticas (conexao com banco, transacoes de reserva).",
    ]
    for item in manut_err:
        add_bullet(doc, item)

    add_para(doc, "Testes e Qualidade", bold=True, size=11)
    manut_test = [
        "Checklist manual de testes por funcionalidade (planilha com cenarios, resultado esperado e resultado obtido).",
        "Testes de integracao: verificar que a constraint UNIQUE(id_sala, data, turno) realmente impede conflitos.",
        "Testes de permissao: tentar cada operacao com cada perfil e confirmar que RBAC funciona.",
        "Testes de borda: curso que cruza virada de ano, turma so de sabado, sala com 3 turnos lotados.",
        "Para o futuro (v2.0+): PHPUnit para testes automatizados dos Models. Nao e viavel no prazo do curso tecnico, mas a arquitetura MVC facilita a adocao posterior.",
    ]
    for item in manut_test:
        add_bullet(doc, item)

    add_para(doc, "Documentacao Tecnica", bold=True, size=11)
    manut_doc = [
        "README.md com instrucoes de instalacao, configuracao e deploy.",
        "CLAUDE.md com contexto do projeto para ferramentas de IA.",
        "Schema SQL comentado (database/schema.sql).",
        "Estrutura de pastas autodocumentada: nomes claros, hierarquia logica.",
        "Sem comentarios obvios no codigo. Comentar apenas o 'por que', nunca o 'o que'.",
    ]
    for item in manut_doc:
        add_bullet(doc, item)

    p_nota_manut = add_para(doc,
        "Nota do desenvolvedor: O documento anterior do grupo mencionava CI/CD, Sentry, ELK Stack e 70% de cobertura de testes unitarios. Essas ferramentas sao excelentes em projetos profissionais, mas estao fora do escopo de um projeto academico hospedado no InfinityFree. As praticas acima sao o que um time junior consegue implementar de forma realista e ja representam um nivel de organizacao acima da media para projetos de curso tecnico. A arquitetura MVC limpa e a separacao de responsabilidades sao a base que torna possivel adotar essas ferramentas no futuro sem reescrever o sistema.",
        italic=True, size=10, color=CINZA, space_after=12)

    doc.add_paragraph()

    # =========================================================
    # 8. REGRAS DE NEGOCIO
    # =========================================================
    add_heading_styled(doc, "8. Regras de Negocio", level=1, color=LARANJA)

    rns = [
        ("RN01", "Reserva por turno, nao por horario livre",
         "A unidade minima de ocupacao e: sala + dia + turno (Manha, Tarde ou Noite). Nao existe reserva por horario livre. Uma sala pode ter ate 3 turmas diferentes no mesmo dia (uma por turno)."),
        ("RN02", "Proibido conflito de sala + turno + dia",
         "O sistema nao deve permitir que duas turmas sejam alocadas na mesma sala, no mesmo turno, no mesmo dia. Implementado como UNIQUE(id_sala, data, turno). Exibir mensagem clara com alternativas."),
        ("RN03", "Dia sem aula estende o termino do curso",
         "Cancelamento adiciona +1 dia util ao final, pulando feriados/recessos e dias fora do padrao. Progresso reflete apenas dias efetivamente realizados."),
        ("RN04", "Cada unidade define seus proprios feriados",
         "Feriados e recessos sao independentes por unidade do SENAC. Preparado para multi-unidade."),
        ("RN05", "Sala em manutencao bloqueia novas reservas",
         "Status Manutencao impede novas reservas. Reservas futuras existentes geram alerta para realocacao. Apenas Suporte pode alterar."),
        ("RN06", "Cursos podem cruzar a virada do ano",
         "Um curso de out/2026 pode terminar em mar/2027. O sistema trata datas reais, sem limitacao a um ano."),
    ]
    for rn_id, rn_title, rn_desc in rns:
        p = doc.add_paragraph()
        r = p.add_run(f"{rn_id} — {rn_title}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = LARANJA
        add_para(doc, rn_desc, size=11)

    # =========================================================
    # 9. CENARIOS DE EXCECAO
    # =========================================================
    add_heading_styled(doc, "9. Cenarios de Excecao", level=1)

    cenarios = [
        ("CE01 — Conflito de reserva",
         "Suporte tenta reservar Sala 5 para Turma T02, turno Manha, mas T01 ja ocupa.",
         "Todas as turmas com reservas futuras na Sala 5.",
         "Bloqueia. Exibe salas disponiveis no mesmo turno/dia. Sugere alternativa.",
         "'A Sala 5 ja esta reservada para a Turma T01 neste turno. Salas disponiveis: Sala 3, Sala 7, Lab 2.'",
         "Nenhuma alteracao no banco. Reserva existente permanece."),
        ("CE02 — Sala entra em manutencao com reservas",
         "Professor reporta defeito. Suporte muda status para Manutencao, mas ha 12 reservas futuras.",
         "Turmas com reservas futuras na sala.",
         "Lista reservas afetadas. Solicita confirmacao antes de alterar.",
         "'Existem 12 reservas futuras nesta sala. Deseja prosseguir?'",
         "Status muda. Reservas permanecem com alerta visual. Log registrado."),
        ("CE03 — Feriado em dia com reservas",
         "Supervisao cadastra feriado 21/04, mas ha 5 reservas nesse dia.",
         "Turmas com aulas previstas no dia.",
         "Alerta. Pergunta se cancela automaticamente. Recalcula terminos (+1 dia).",
         "'Existem 5 reservas no dia 21/04. Deseja cancelar e recalcular?'",
         "Feriado cadastrado. Reservas canceladas. Terminos recalculados."),
        ("CE04 — Curso cruza virada de ano",
         "Turma 01/10/2026 a 15/02/2027. Reservas recorrentes precisam cruzar o ano.",
         "Suporte ao gerar reservas.",
         "Gera reservas normalmente em ambos os anos, respeitando feriados.",
         "'Foram geradas X reservas de 01/10/2026 a 15/02/2027, pulando Y feriados.'",
         "Reservas criadas em ambos os anos. Progresso calculado corretamente."),
        ("CE05 — Acesso nao autorizado",
         "Docente tenta acessar URL de criacao de reserva diretamente.",
         "Docente (acao impedida).",
         "Verifica perfil. Nega acesso. Redireciona ao painel.",
         "'Voce nao tem permissao para acessar esta funcionalidade.'",
         "Nenhuma alteracao. Tentativa registrada no log."),
    ]

    for ce_title, aconteceu, afetado, acao, msg, dados in cenarios:
        p = doc.add_paragraph()
        r = p.add_run(ce_title)
        r.bold = True
        r.font.size = Pt(11)
        r.font.color.rgb = AZUL
        add_para(doc, f"O que aconteceu: {aconteceu}", size=10)
        add_para(doc, f"Quem foi afetado: {afetado}", size=10)
        add_para(doc, f"Acao do sistema: {acao}", size=10)
        add_para(doc, f"Mensagem: {msg}", italic=True, size=10)
        add_para(doc, f"Dados: {dados}", size=10)

    p_ce = doc.add_paragraph()
    r_ce = p_ce.add_run("(Cenarios detalhados conforme roteiro do professor)")
    r_ce.italic = True
    r_ce.font.size = Pt(10)
    r_ce.font.color.rgb = CINZA
    add_tracked_comment(p_ce,
        "MINHA CONTRIBUICAO: Criei todos os 5 cenarios de excecao seguindo o modelo do roteiro do professor (o que aconteceu, quem foi afetado, acao do sistema, mensagem, dados). O documento do grupo nao inclui cenarios de excecao.",
        all_comments)

    # =========================================================
    # 10. MOSCOW
    # =========================================================
    add_heading_styled(doc, "10. Priorizacao MoSCoW", level=1)
    add_para(doc, "Criterio: o que e necessario para substituir a planilha sem perder funcionalidade. Must Have = 60% do esforco (regra DSDM).")

    moscow_table = doc.add_table(rows=1, cols=4)
    moscow_table.style = 'Table Grid'
    hdr = moscow_table.rows[0]
    moscow_headers = [
        ("Must — Obrigatorio\n(Sem isso nao funciona)", "C62828"),
        ("Should — Importante\n(MVP pode lancar sem)", "E65100"),
        ("Could — Desejavel\n(Se der tempo)", "2E7D32"),
        ("Won't — Futuro\n(Fora do escopo v1)", "757575"),
    ]
    for i, (txt, cor) in enumerate(moscow_headers):
        hdr.cells[i].text = txt
        set_cell_shading(hdr.cells[i], cor)
        for p in hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.bold = True
                r.font.size = Pt(9)

    row = moscow_table.add_row()
    must_items = "RF01. Autenticacao\nRF02. Cadastro de salas\nRF04. Cadastro de turmas\nRF06. Reserva + conflito\nRF10. Status da sala\nRF11. Painel/calendario\nRF12. Filtros\nRN01-02. Turno + constraint"
    should_items = "RF03. Cadastro cursos\nRF05. Feriados/recessos\nRF07. Reserva recorrente\nRF09. Dia sem aula\nRF13. Historico por turma\nRF14. Progresso\nRF15. Gerenciar docentes"
    could_items = "RF08. Realocar turma\nRF16. Log auditoria\nRF17. Dashboard\nImpressao/PDF"
    wont_items = "Integracao chamados\nNotificacoes email/SMS\nApp mobile nativo\nSugestao automatica\nMulti-unidade ativo\nExportacao CSV/Excel\nAlerta fim de turma"

    for i, items in enumerate([must_items, should_items, could_items, wont_items]):
        row.cells[i].text = items
        for p in row.cells[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)

    doc.add_paragraph()

    # =========================================================
    # 11. GLOSSARIO
    # =========================================================
    add_heading_styled(doc, "11. Glossario / Dicionario de Dados", level=1)

    glossario = [
        ("Sala", "Ambiente fisico. Tipos: Sala de Aula, Laboratorio, Auditorio, Cozinha, Especial."),
        ("Turma", "Grupo de alunos com codigo unico (006.2026.XXXX). Tem turno, professor(es) e dias fixos."),
        ("Reserva", "Alocacao: sala + dia + turno. Unidade minima de ocupacao."),
        ("Turno", "Manha, Tarde ou Noite. Ate 3 turmas por dia na mesma sala."),
        ("Andar", "Pavimento do predio. Confirmado pela cliente como importante. Cadastro livre."),
        ("Progresso", "(dias realizados / dias previstos) x 100%. Calculado automaticamente."),
        ("Dia cancelado", "Dia previsto sem aula. Adiciona +1 dia util ao termino."),
        ("Status da sala", "Disponivel, Ocupada ou Manutencao."),
        ("Unidade", "Filial SENAC. v1: somente BH. Preparado para expansao."),
        ("Padrao de dias", "Dias da semana da turma (Seg-Sex, Sab, Seg/Qua/Sex)."),
        ("Tipo de Curso", "Cursos Ageis, Graduacao, Pos Graduacao, Tecnico."),
        ("Area de Conhecimento", "23 areas SENAC (Gestao, TI, Gastronomia, Design, Saude, Games, etc.)."),
        ("Programa", "Todos os cursos, Senac+, PSG (Programa SENAC de Gratuidade)."),
        ("RBAC", "Controle de acesso por perfil (Admin, Supervisao, Suporte, Docente)."),
        ("InfinityFree", "Hospedagem gratuita PHP + MySQL. Sem SSH. SSL via Let's Encrypt."),
    ]

    g_table = doc.add_table(rows=1, cols=2)
    g_table.style = 'Table Grid'
    g_hdr = g_table.rows[0]
    for i, txt in enumerate(["Termo", "Definicao"]):
        g_hdr.cells[i].text = txt
        set_cell_shading(g_hdr.cells[i], "004C94")
        for p in g_hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.bold = True
                r.font.size = Pt(10)

    for termo, definicao in glossario:
        row = g_table.add_row()
        row.cells[0].text = termo
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = definicao
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================
    # 12. ROTEIRO DE PERGUNTAS
    # =========================================================
    add_heading_styled(doc, "12. Roteiro de Perguntas-Chave", level=1)

    perguntas = [
        ("1", "Quando troca de sala, quem decide?", "Suporte decide. Comunicacao informal. Futuro: Admin eleito."),
        ("2", "Ja teve conflito de sala?", "Sim. Supervisao detecta visualmente e resolve."),
        ("3", "Mais alguem acessa?", "Sim. Supervisao, Suporte e Docentes."),
        ("4", "Salas entram em manutencao?", "Sim. Chamado no canal oficial, Suporte altera status."),
        ("5", "Info que mais demora?", "Historico de salas por turma."),
        ("6", "Progresso automatico?", "Sim. Dia sem aula = +1 dia no termino."),
        ("7", "Sabado letivo?", "Nao para todos. Turmas Seg-Sex e turmas so Sabado."),
        ("8", "Tem andares?", "Sim. Andares diferentes. Numero importante. Cadastro livre."),
        ("9", "Mais de um professor?", "Sim, varios professores por turma."),
        ("10", "Feriados proprios?", "Cada unidade define os seus."),
    ]

    q_table = doc.add_table(rows=1, cols=3)
    q_table.style = 'Table Grid'
    q_hdr = q_table.rows[0]
    for i, txt in enumerate(["#", "Pergunta", "Resposta"]):
        q_hdr.cells[i].text = txt
        set_cell_shading(q_hdr.cells[i], "004C94")
        for p in q_hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.bold = True
                r.font.size = Pt(10)

    for num, pergunta, resposta in perguntas:
        row = q_table.add_row()
        row.cells[0].text = num
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = pergunta
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[2].text = resposta
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================
    # 13. ARTEFATOS
    # =========================================================
    add_heading_styled(doc, "13. Artefatos de Modelagem e Analise", level=1)

    artefatos = [
        ("DER", "11 tabelas + 4 views com tipos, PKs, FKs", "docs/02-analise/DER_Mapa_de_Salas.html"),
        ("Schema SQL", "DDL completo do banco", "database/schema.sql"),
        ("Prototipos", "7 telas com paleta SENAC", "docs/04-design/Prototipos_Telas.html"),
        ("PRD", "Produto, cronograma, riscos", "docs/02-analise/PRD_Mapa_de_Salas.html"),
        ("Roteiro", "Preenchido em sala pelo professor", "docs/01-levantamento/Roteiro_Analise_Requisitos.html"),
    ]

    a_table = doc.add_table(rows=1, cols=3)
    a_table.style = 'Table Grid'
    a_hdr = a_table.rows[0]
    for i, txt in enumerate(["Artefato", "Descricao", "Arquivo"]):
        a_hdr.cells[i].text = txt
        set_cell_shading(a_hdr.cells[i], "004C94")
        for p in a_hdr.cells[i].paragraphs:
            for r in p.runs:
                r.font.color.rgb = BRANCO
                r.bold = True
                r.font.size = Pt(10)

    for art, desc, arq in artefatos:
        row = a_table.add_row()
        row.cells[0].text = art
        row.cells[0].paragraphs[0].runs[0].bold = True
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].text = desc
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[2].text = arq
        row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)

    doc.add_paragraph()

    # =========================================================
    # 14. VALIDACAO
    # =========================================================
    add_heading_styled(doc, "14. Processo de Validacao e Aprovacao", level=1)

    validacao = [
        "Review Interno: Equipe revisa o documento para garantir coerencia entre requisitos, regras de negocio e schema do banco.",
        "Apresentacao (Walkthrough): Prototipos, DER e requisitos priorizados para a supervisao/cliente.",
        "Ajustes: Coletar feedback, corrigir interpretacoes, adicionar requisitos faltantes.",
        "Aprovacao: Cliente aprova. Mudancas de escopo = renegociacao de prazo.",
    ]
    for i, v in enumerate(validacao, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"{i}. ")
        r.bold = True
        r.font.size = Pt(11)
        r2 = p.add_run(v)
        r2.font.size = Pt(11)

    p_unif = add_para(doc,
        "Unificacao (v3.0): Esta versao consolida as analises de toda a equipe em um unico documento. Correcoes aplicadas: perfis alinhados com a entrevista (Supervisao/Suporte/Docente/Admin), reserva por turno, multiplos professores (N:N), andar obrigatorio, tipos de curso e 23 areas SENAC, RNFs realistas para InfinityFree. Removidos do documento anterior: SLA 99,5%, Docker, CI/CD, Sentry, ELK Stack, 70% cobertura de testes (incompativeis com o escopo academico).",
        italic=True, size=10, color=CINZA)

    # =========================================================
    # 15. CHECKLIST
    # =========================================================
    add_heading_styled(doc, "15. Checklist de Encerramento da Etapa", level=1)

    checklist = [
        (True, "Todos os stakeholders-chave foram entrevistados?"),
        (True, "Os requisitos funcionais estao escritos de forma clara e testavel?"),
        (True, "Os requisitos nao funcionais estao definidos com metricas realistas?"),
        (True, "As regras de negocio estao documentadas separadamente?"),
        (True, "Os cenarios de excecao estao detalhados?"),
        (True, "As prioridades MoSCoW estao definidas?"),
        (True, "O glossario contem todos os termos do dominio?"),
        (True, "Os prototipos de interface foram criados?"),
        (True, "O DER e schema SQL estao alinhados com os requisitos?"),
        (False, "O documento foi aprovado e assinado pelo patrocinador do projeto?"),
    ]
    for done, item in checklist:
        mark = "[x]" if done else "[ ]"
        p = doc.add_paragraph()
        r = p.add_run(f"  {mark}  {item}")
        r.font.size = Pt(11)
        if done:
            r.font.color.rgb = VERDE_ESCURO
        else:
            r.font.color.rgb = VERMELHO

    doc.add_paragraph()
    p_dica = add_para(doc,
        'Dica de Ouro (guia do professor): "Nunca escreva requisitos usando \'etc\', \'entre outros\' ou \'diversos\'. Se voce colocou isso no documento, significa que a coleta nao terminou. Seja especifico."',
        italic=True, size=10, color=LARANJA)

    # =========================================================
    # RODAPE
    # =========================================================
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Documento de Requisitos Unificado — Mapa de Salas SENAC Minas BH | Versao 3.0 | Julho 2026")
    r.font.size = Pt(10)
    r.font.color.rgb = CINZA
    footer.add_run("\n")
    r2 = footer.add_run("Elaborado por Thais Oliveira e equipe | Tecnico em Desenvolvimento de Sistemas — SENAC Minas")
    r2.font.size = Pt(10)
    r2.font.color.rgb = CINZA
    footer.add_run("\n")
    r3 = footer.add_run("Modelo conforme guia do professor Cleiton de Jesus Pereira | Hospedagem MVP: InfinityFree")
    r3.font.size = Pt(10)
    r3.font.color.rgb = CINZA

    # =========================================================
    # FINALIZAR COMENTARIOS E SALVAR
    # =========================================================
    if all_comments:
        finalize_comments(doc, all_comments)

    output_path = r"C:\Users\thais\Downloads\projeto_mapa_de_sala_senac\docs\01-levantamento\Documento_de_Requisitos_Unificado_v3.docx"
    doc.save(output_path)
    print(f"Documento salvo em: {output_path}")
    print(f"Total de comentarios: {len(all_comments)}")
    for i, (cid, author, initials, text) in enumerate(all_comments, 1):
        print(f"  Comentario {i}: {text[:80]}...")


if __name__ == "__main__":
    main()
